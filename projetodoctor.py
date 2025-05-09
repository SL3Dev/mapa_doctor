import streamlit as st
from graphviz import Digraph
from PyPDF2 import PdfReader
from docx import Document
import tempfile
import os
import re
from graphviz import backend

# --- Verificação de Instalação do Graphviz --- #
def check_graphviz_installed():
    try:
        backend.execute.run_check(['dot', '-V'])
        return True
    except backend.ExecutableNotFound:
        return False

GRAPHVIZ_INSTALLED = check_graphviz_installed()

# --- Dados Médicos Aprimorados com Ícones --- #
RELACOES_MEDICAS = {
    "🦠 Lesão Celular": {
        "📌 pode_ser": ["🔄 Reversível", "❌ Irreversível"],
        "⚡ causas": ["🫁 Hipóxia", "☠️ Toxinas", "🦠 Infecções"],
        "📚 exemplo": ["Paciente com hepatite alcoólica"]
    },
    "🫁 Hipóxia": {
        "🔬 causa": ["🔥 Estresse Oxidativo (ROS)", "🧬 Dano Mitocondrial"],
        "⚠️ leva_a": ["💀 Necrose", "⚰️ Apoptose"],
        "📚 exemplo": ["Isquemia miocárdica em IAM"]
    },
    "🔥 Inflamação": {
        "👀 caracteriza_se_por": ["🩸 Vasodilatação", "🔴 Eritema", "💧 Edema"],
        "📌 pode_levar": ["🩹 Reparo Tecidual", "🧶 Fibrose"],
        "📚 exemplo": ["Apendicite aguda"]
    },
    "💀 Necrose": {
        "📝 tipos": ["🪨 Coagulativa", "💧 Liquefativa", "🧀 Caseosa"],
        "📚 exemplo": ["Gangrena em membro diabético"]
    },
    "⚰️ Apoptose": {
        "🔍 diferenciacao": ["⏱️ Morte Programada", "🚫 Sem Inflamação"],
        "📚 exemplo": ["Involução uterina pós-parto"]
    }
}

# --- Cores e Estilos Aprimorados --- #
CORES = {
    "causa": "#FF6B6B",  # Vermelho
    "sintoma": "#4ECDC4",  # Azul
    "tratamento": "#45B7D1",  # Azul forte
    "diagnóstico": "#FFA07A",  # Laranja
    "prevenção": "#7BC043",  # Verde
    "exemplo": "#A593E0"  # Roxo
}

def get_cor_relacao(relacao):
    relacao = relacao.lower()
    if any(p in relacao for p in ["causa", "leva_a"]): return CORES["causa"]
    elif any(p in relacao for p in ["sintoma", "caracteriza"]): return CORES["sintoma"]
    elif "trata" in relacao: return CORES["tratamento"]
    elif "diagnóstico" in relacao: return CORES["diagnóstico"]
    elif "previne" in relacao: return CORES["prevenção"]
    elif "exemplo" in relacao: return CORES["exemplo"]
    return "#6B5B95"  # Cor padrão

# --- Funções Aprimoradas --- #
def processar_texto(texto):
    """Normaliza o texto para análise"""
    texto = re.sub(r'[^\w\sáéíóúâêîôûãõàèìòùäëïöüç]', ' ', texto, flags=re.IGNORECASE)
    texto = re.sub(r'\s+', ' ', texto).strip()
    return texto.lower()

def extrair_texto(file):
    """Extrai texto de arquivos com tratamento de erro melhorado"""
    try:
        if file.type == "application/pdf":
            pdf = PdfReader(file)
            texto = " ".join(page.extract_text() or "" for page in pdf.pages)
            return processar_texto(texto)
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(file)
            texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return processar_texto(texto)
        elif file.type == "text/plain":
            texto = file.read().decode("utf-8")
            return processar_texto(texto)
        return ""
    except Exception as e:
        st.error(f"Erro ao extrair texto: {str(e)}")
        return ""

def gerar_proposicoes_offline(texto):
    """Gera relações médicas a partir do texto de entrada de forma dinâmica"""
    texto = texto.lower()
    proposicoes = []
    termos_encontrados = False
    
    # Dicionário de mapeamento de termos sem emojis
    termos_mapeamento = {
        "lesão celular": "🦠 Lesão Celular",
        "lesao celular": "🦠 Lesão Celular",
        "hipóxia": "🫁 Hipóxia",
        "hipoxia": "🫁 Hipóxia",
        "inflamação": "🔥 Inflamação",
        "inflamacao": "🔥 Inflamação",
        "necrose": "💀 Necrose",
        "apoptose": "⚰️ Apoptose"
    }
    
    # Verifica cada termo médico no texto
    for termo_sem_icone, termo_com_icone in termos_mapeamento.items():
        if termo_sem_icone in texto:
            termos_encontrados = True
            # Adiciona todas as relações para este termo
            for relacao, termos in RELACOES_MEDICAS[termo_com_icone].items():
                for termo in termos:
                    proposicoes.append(f"{termo_com_icone} -> {relacao} -> {termo}")
    
    # Se não encontrou termos médicos, usa exemplos pré-definidos
    if not termos_encontrados:
        st.warning("Nenhum termo médico relevante detectado. Mostrando exemplos genéricos.")
        for termo_principal in RELACOES_MEDICAS:
            for relacao, termos in RELACOES_MEDICAS[termo_principal].items():
                for termo in termos:
                    proposicoes.append(f"{termo_principal} -> {relacao} -> {termo}")
    
    return "\n".join(proposicoes)

def criar_mapa_avancado(proposicoes, orientacao="retrato"):
    """Cria mapa conceitual com ícones e estilo aprimorado"""
    rankdir = "TB" if orientacao == "retrato" else "LR"
    
    grafo = Digraph(
        graph_attr={
            "rankdir": rankdir,
            "bgcolor": "#f9f9f9",
            "fontname": "Arial",
            "splines": "ortho",
            "compound": "true"
        },
        node_attr={
            "shape": "rectangle",
            "style": "filled,rounded",
            "fillcolor": "#ffffff",
            "fontname": "Arial",
            "fontsize": "12",
            "penwidth": "1.5"
        },
        edge_attr={
            "fontname": "Arial",
            "fontsize": "10",
            "penwidth": "1.2"
        }
    )

    if not proposicoes:
        return grafo

    conceitos = {}
    for linha in proposicoes.split("\n"):
        if "->" in linha:
            partes = [p.strip() for p in linha.split("->")]
            if len(partes) >= 3:
                origem, relacao, destino = partes[0], partes[1], "->".join(partes[2:])
                if origem not in conceitos:
                    conceitos[origem] = []
                conceitos[origem].append((relacao, destino))

    for conceito, relacoes in conceitos.items():
        with grafo.subgraph(name=f"cluster_{conceito}") as sub:
            sub.attr(
                style="filled",
                color="#e0e0e0",
                fillcolor="#f5f5f5",
                label=conceito,
                fontsize="12",
                fontcolor="#333333"
            )
            
            sub.node(conceito, 
                    shape="ellipse", 
                    fillcolor="#e6f3ff", 
                    style="filled",
                    penwidth="2.0")
            
            for relacao, destino in relacoes:
                cor = get_cor_relacao(relacao)
                
                if "exemplo" in relacao:
                    sub.node(destino, 
                            shape="note", 
                            fillcolor="#f0e6ff",
                            style="filled")
                else:
                    sub.node(destino)
                
                sub.edge(conceito, destino, 
                        label=relacao, 
                        color=cor, 
                        fontcolor=cor,
                        penwidth="1.5")

    return grafo

def exportar_para_pdf(grafo, orientacao="retrato"):
    """Exporta para PDF com metadados médicos"""
    rankdir = "TB" if orientacao == "retrato" else "LR"
    grafo.graph_attr["rankdir"] = rankdir
    grafo.graph_attr["label"] = "Mapa Conceitual Médico\n\n"
    grafo.graph_attr["labelloc"] = "t"
    grafo.graph_attr["fontsize"] = "16"
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmpfile:
        temp_path = tmpfile.name
    
    try:
        grafo.render(temp_path.replace('.pdf', ''), format='pdf', cleanup=True)
        with open(temp_path, 'rb') as f:
            pdf_data = f.read()
        return pdf_data
    finally:
        if os.path.exists(temp_path):
            os.unlink(temp_path)
        if os.path.exists(temp_path.replace('.pdf', '')):
            os.unlink(temp_path.replace('.pdf', ''))

def exportar_para_word(grafo, orientacao="retrato"):
    """Exporta para Word com formatação médica"""
    rankdir = "TB" if orientacao == "retrato" else "LR"
    grafo.graph_attr["rankdir"] = rankdir
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as imgfile:
        img_path = imgfile.name
    
    try:
        grafo.render(img_path.replace('.png', ''), format='png', cleanup=True)
        
        if not os.path.exists(img_path):
            raise FileNotFoundError("Falha ao gerar imagem do gráfico")
        
        doc = Document()
        doc.add_heading('Mapa Conceitual Médico', level=1)
        doc.add_paragraph('Gerado por:', style='Heading 2')
        doc.add_paragraph('Sistema de Mapeamento Conceitual Médico v2.0')
        doc.add_paragraph(f'Orientação: {orientacao.capitalize()}')
        doc.add_picture(img_path, width=docx.shared.Inches(6))
        doc.add_paragraph().add_run('Referências:').bold = True
        doc.add_paragraph('- Robbins & Cotran: Bases Patológicas das Doenças')
        doc.add_paragraph('- Guyton & Hall: Tratado de Fisiologia Médica')
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmpfile:
            temp_path = tmpfile.name
            doc.save(temp_path)
        
        with open(temp_path, 'rb') as f:
            word_data = f.read()
        return word_data
    finally:
        if os.path.exists(img_path):
            os.unlink(img_path)
        if os.path.exists(img_path.replace('.png', '')):
            os.unlink(img_path.replace('.png', ''))
        if os.path.exists(temp_path):
            os.unlink(temp_path)

def main():
    st.set_page_config(
        page_title="Mapa Conceitual Médico Avançado",
        page_icon="🩺",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    with st.sidebar:
        st.title("🩺 Ajuda Médica")
        st.markdown("""
        **Legenda de Ícones:**
        - 🦠: Processos patológicos
        - 🫁: Condições fisiológicas
        - 🔬: Mecanismos
        - ⚠️: Complicações
        - 📚: Exemplos clínicos
        """)
        
        if not GRAPHVIZ_INSTALLED:
            st.warning("""
            **Graphviz não instalado!**
            Para exportar PDF/Word:
            1. [Baixe o Graphviz](https://graphviz.org/download/)
            2. Instale com opção "Add to PATH"
            3. Reinicie o aplicativo
            """)

    st.title("🧠 Mapa Conceitual Médico Avançado")
    st.markdown("""
    *Ferramenta para organização visual de conceitos médicos*  
    **Como usar**:  
    1. Insira texto ou faça upload de arquivos  
    2. Ajuste as configurações de visualização  
    3. Explore as relações conceituais  
    4. Exporte para estudo ou apresentação  
    """)

    tab1, tab2 = st.tabs(["✍️ Digitar Texto", "📂 Upload de Arquivos"])
    
    with tab1:
        texto_manual = st.text_area("Cole texto médico aqui:", height=300,
                                  placeholder="Ex: Lesões celulares podem ser reversíveis (esteatose) ou irreversíveis (necrose)...",
                                  key="texto_input")
    
    with tab2:
        arquivos = st.file_uploader("Selecione arquivos (PDF, DOCX, TXT)", 
                                  type=["pdf", "docx", "txt"],
                                  accept_multiple_files=True,
                                  key="file_uploader")

    with st.expander("⚙️ Configurações Avançadas"):
        col1, col2 = st.columns(2)
        with col1:
            orientacao = st.radio("Orientação:", ["Retrato", "Paisagem"], index=0, key="orientacao")
        with col2:
            if GRAPHVIZ_INSTALLED:
                layout = st.selectbox(
                    "Algoritmo de Layout:",
                    ["dot", "neato", "fdp", "sfdp", "twopi", "circo"],
                    index=0,
                    key="layout"
                )

    if st.button("🔄 Gerar Mapa Conceitual", type="primary", use_container_width=True, key="gerar_mapa"):
        conteudo_total = ""
        if texto_manual:
            conteudo_total += processar_texto(texto_manual) + "\n"
        
        if arquivos:
            with st.spinner("Processando arquivos..."):
                for arquivo in arquivos:
                    conteudo = extrair_texto(arquivo)
                    if conteudo:
                        conteudo_total += conteudo + "\n"
        
        if not conteudo_total.strip():
            st.warning("Por favor, insira texto ou faça upload de arquivos!")
            st.stop()
        
        with st.spinner("Identificando relações médicas..."):
            proposicoes = gerar_proposicoes_offline(conteudo_total)
        
        st.success("Análise concluída! Visualizando relações...")
        
        with st.expander("📋 Relações Identificadas"):
            st.code(proposicoes)
        
        st.subheader("🧬 Visualização do Mapa Conceitual")
        grafo = criar_mapa_avancado(proposicoes, orientacao.lower())
        
        if GRAPHVIZ_INSTALLED:
            grafo.engine = layout
        
        st.graphviz_chart(grafo, use_container_width=True)
        
        st.subheader("💾 Opções de Exportação")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.download_button(
                label="📝 Baixar como DOT",
                data=grafo.source,
                file_name="mapa_medico.dot",
                mime="text/plain",
                help="Formato Graphviz para edição avançada",
                use_container_width=True
            )
        
        with col2:
            if GRAPHVIZ_INSTALLED:
                pdf_data = exportar_para_pdf(grafo, orientacao.lower())
                st.download_button(
                    label="📄 Baixar como PDF",
                    data=pdf_data,
                    file_name="mapa_medico.pdf",
                    mime="application/pdf",
                    help="Documento PDF pronto para impressão",
                    use_container_width=True
                )
            else:
                st.button(
                    "📄 Baixar como PDF (Graphviz necessário)",
                    disabled=True,
                    help="Instale o Graphviz para habilitar esta função",
                    use_container_width=True
                )
        
        with col3:
            if GRAPHVIZ_INSTALLED:
                word_data = exportar_para_word(grafo, orientacao.lower())
                st.download_button(
                    label="📑 Baixar como Word",
                    data=word_data,
                    file_name="mapa_medico.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    help="Documento editável com formatação profissional",
                    use_container_width=True
                )
            else:
                st.button(
                    "📑 Baixar como Word (Graphviz necessário)",
                    disabled=True,
                    help="Instale o Graphviz para habilitar esta função",
                    use_container_width=True
                )

    with st.expander("🧪 Exemplos Clínicos Prontos"):
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("Caso 1: Lesão Hepática", key="exemplo1"):
                st.session_state.texto_exemplo = """
                Paciente masculino, 45 anos, etilista crônico apresenta:
                - 🦠 Lesão celular hepática 🔄 Reversível (esteatose)
                - ☠️ Toxinas: Etanol
                - 🩹 Reparo tecidual com abstinência
                """
        
        with col2:
            if st.button("Caso 2: Infarto do Miocárdio", key="exemplo2"):
                st.session_state.texto_exemplo = """
                Paciente feminino, 60 anos, diabética:
                - 🫁 Hipóxia miocárdica por aterosclerose
                - 🔥 Estresse oxidativo → 💀 Necrose 🪨 Coagulativa
                - ⚠️ Complicação: Fibrose cardíaca
                """
        
        if 'texto_exemplo' in st.session_state:
            texto_manual = st.text_area("Texto de exemplo:", 
                                      value=st.session_state.texto_exemplo, 
                                      height=150,
                                      key="texto_exemplo_area")

if __name__ == "__main__":
    main()